import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_reconstructed_document(abstract_text: str, insights, image_paths: list, output_filename: str = "Reconstructed_Paper.docx"):
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI RECONSTRUCTION REPORT")
    title_run.font.size = Pt(18)
    title_run.bold = True
    
    doc.add_paragraph("\n")
    
    # Section 1: Core Insights Table
    doc.add_heading("1. Executive Technical Summary", level=1)
    doc.add_paragraph(f"Core Scientific Contribution: {insights.core_contribution}")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technical Dimension'
    hdr_cells[1].text = 'Extracted Detail'
    
    table.add_row().cells[0].text, table.rows[1].cells[1].text = 'Architecture', insights.architecture
    table.add_row().cells[0].text, table.rows[2].cells[1].text = 'Datasets', ", ".join(insights.datasets)
    
    metrics_str = "\n".join([f"{k}: {v}" for k, v in insights.key_metrics.items()])
    table.add_row().cells[0].text, table.rows[3].cells[1].text = 'Key Metrics', metrics_str

    # Section 2: Reconstructed Abstract
    doc.add_heading("2. Reconstructed Abstract", level=1)
    p_abstract = doc.add_paragraph(abstract_text)
    p_abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Section 3: Key Calculations & Formulas
    doc.add_heading("3. Core Mathematical Calculations", level=1)
    if insights.key_calculations:
        for calc in insights.key_calculations:
            doc.add_paragraph(calc, style='List Bullet')
    else:
        doc.add_paragraph("No significant mathematical formulas detected in the parsed text.")
        
    # Section 4: Extracted Graphs and Visuals
    doc.add_heading("4. Extracted Graphs & Visuals", level=1)
    if image_paths:
        for img_path in image_paths:
            try:
                # Insert the image and scale it to fit the page
                doc.add_picture(img_path, width=Inches(5.5))
                doc.add_paragraph(f"Source file: {os.path.basename(img_path)}")
                doc.add_paragraph("\n")
            except Exception as e:
                print(f"Failed to insert image {img_path}: {e}")
    else:
        doc.add_paragraph("No charts or graphs of significant size were found in this paper.")

    # Section 5: Comprehensive Final Summary
    doc.add_heading("5. Final Holistic Summary", level=1)
    p_summary = doc.add_paragraph(insights.final_summary)
    p_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "data")
    file_path = os.path.join(output_dir, output_filename)
    doc.save(file_path)
    return file_path