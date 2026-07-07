"""
frontend/app.py
----------------
Streamlit UI for the Research Reconstructor pipeline.

Responsibilities:
  1. Let the user upload a research-paper PDF.
  2. Split the PDF into text chunks (for local Ollama graph extraction) and
     raw chart images (for Groq vision analysis).
  3. Kick off the hybrid backend pipeline:
        - backend.services.graph_store  -> local Ollama + Neo4j Aura
        - backend.services.vision       -> PIL compression + Groq vision
  4. Ask a Groq-hosted model to synthesize a 3-4 paragraph technical
     executive summary from the structural + visual insights.
  5. Compile everything into a downloadable .docx via python-docx.

Historical bottleneck resolved here (Section 5 of master context):
  The .docx compiler used to crash with `'str' object has no attribute
  'items'` whenever it hit a chart entry the vision API had skipped/returned
  as a plain fallback string. The compiler loop below explicitly checks
  `isinstance(data, dict)` before treating an entry as a structured result,
  so a bare string (or any unexpected type) degrades gracefully instead of
  crashing the whole document build.
"""

import os
import sys
import io
import logging
import tempfile
from datetime import datetime
from typing import List, Dict

import streamlit as st
import fitz  # PyMuPDF
import httpx
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from dotenv import load_dotenv
from groq import Groq, APIError, RateLimitError, APIConnectionError

# Make backend/ importable when running `streamlit run frontend/app.py`
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from backend.services.graph_store import GraphStore
from backend.services.vision import extract_images_from_pdf, analyze_chart

load_dotenv()

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_TEXT_MODEL = os.getenv("GROQ_TEXT_MODEL", "openai/gpt-oss-120b")
GROQ_TIMEOUT_SECONDS = float(os.getenv("GROQ_TIMEOUT_SECONDS", "180"))

CHUNK_SIZE_CHARS = 3000
CHUNK_OVERLAP_CHARS = 200


# ---------------------------------------------------------------------------
# Groq client (synthesis layer)
# ---------------------------------------------------------------------------
def get_groq_client() -> Groq:
    """
    Same custom-timeout pattern as vision.py — prevents Windows `wsarecv`
    socket aborts during long report-writing generations.
    """
    if not GROQ_API_KEY:
        st.error("GROQ_API_KEY is not set. Please configure your .env file.")
        st.stop()
    http_client = httpx.Client(timeout=GROQ_TIMEOUT_SECONDS)
    return Groq(api_key=GROQ_API_KEY, http_client=http_client)


# ---------------------------------------------------------------------------
# PDF text extraction + chunking
# ---------------------------------------------------------------------------
def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    text_parts = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE_CHARS, overlap: int = CHUNK_OVERLAP_CHARS) -> List[str]:
    """Simple sliding-window character chunker (no external dependency needed)."""
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunks.append(text[start:end])
        start += chunk_size - overlap

    return [c for c in chunks if c.strip()]


# ---------------------------------------------------------------------------
# Report synthesis (Groq)
# ---------------------------------------------------------------------------
def synthesize_report(
    client: Groq,
    paper_text_excerpt: str,
    chart_analyses: List[Dict],
    graph_summary: str,
) -> str:
    """
    Prompts a Groq-hosted model to write a 3-4 paragraph technical executive
    summary grounded in the extracted structural (graph) insights and visual
    (chart) insights.
    """
    chart_notes = "\n".join(
        f"- {c.get('description', 'No description available.')}"
        for c in chart_analyses
        if isinstance(c, dict) and not c.get("skipped", False)
    ) or "No chart-level insights were extracted."

    prompt = f"""
You are a technical research analyst. Write a 3-4 paragraph executive
summary of the following research paper for an informed but time-constrained
reader. Ground your summary in the structural knowledge-graph insights and
the chart insights provided below. Be precise, avoid filler, and highlight
the paper's core contribution, methodology, and key quantitative results.

--- PAPER EXCERPT ---
{paper_text_excerpt[:6000]}

--- KNOWLEDGE GRAPH SUMMARY ---
{graph_summary}

--- CHART INSIGHTS ---
{chart_notes}
""".strip()

    response = None
    try:
        response = client.chat.completions.create(
            model=GROQ_TEXT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=900,
        )
    except RateLimitError as exc:
        # Covers both true rate limiting AND quota/billing errors — Groq's
        # SDK raises the same exception type for both, distinguished only
        # by the error message/code.
        if "insufficient_quota" in str(exc).lower() or "quota" in str(exc).lower():
            raise RuntimeError(
                "Your Groq account has no available quota. Check your usage "
                "limits and billing at console.groq.com/settings/billing, "
                "then try again."
            ) from exc
        raise RuntimeError(
            "Groq rate limit hit. Wait a moment and try again, or check "
            "your plan's requests-per-minute / tokens-per-minute limit at "
            "console.groq.com."
        ) from exc
    except APIConnectionError as exc:
        raise RuntimeError(
            "Could not reach the Groq API. Check your internet connection "
            "and firewall/proxy settings."
        ) from exc
    except APIError as exc:
        raise RuntimeError(f"Groq API error during synthesis: {exc}") from exc

    return response.choices[0].message.content.strip()


def summarize_graph_documents(graph_documents: List) -> str:
    """Turns LangChain GraphDocument objects into a short human-readable summary."""
    if not graph_documents:
        return "No graph entities or relationships were extracted."

    lines = []
    for gdoc in graph_documents:
        node_labels = sorted({node.type for node in getattr(gdoc, "nodes", [])})
        rel_types = sorted({rel.type for rel in getattr(gdoc, "relationships", [])})
        if node_labels:
            lines.append(f"Entity types found: {', '.join(node_labels)}")
        if rel_types:
            lines.append(f"Relationship types found: {', '.join(rel_types)}")

    return "\n".join(lines) if lines else "Graph extraction returned no labeled entities."


# ---------------------------------------------------------------------------
# .docx compiler
# ---------------------------------------------------------------------------
def compile_docx(
    title: str,
    summary_text: str,
    graph_summary: str,
    chart_analyses: List[Dict],
) -> io.BytesIO:
    """
    Builds the final Word document.

    CRITICAL FIX (Section 5 of master context): each entry in chart_analyses
    is checked with `isinstance(entry, dict)` before being treated as a
    structured result. Previously, any non-dict fallback value crashed the
    build with `'str' object has no attribute 'items'`. Now, unexpected
    types are rendered as a generic placeholder line instead of crashing.
    """
    doc = DocxDocument()

    # --- Title ---
    title_heading = doc.add_heading(title, level=0)

    doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    doc.add_paragraph("")

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    for paragraph in summary_text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    # --- Knowledge Graph Insights ---
    doc.add_heading("Knowledge Graph Insights", level=1)
    doc.add_paragraph(graph_summary)

    # --- Chart Analysis ---
    doc.add_heading("Chart & Figure Analysis", level=1)

    if not chart_analyses:
        doc.add_paragraph("No charts or figures were detected in this document.")
    else:
        for idx, entry in enumerate(chart_analyses, start=1):
            # --- MANDATORY type guard: prevents the historical crash ---
            if isinstance(entry, dict):
                if entry.get("skipped", False):
                    text = entry.get("description", "This image was skipped during analysis.")
                else:
                    text = entry.get("description", "No description available.")
            else:
                # Graceful fallback for any unexpected (e.g., string/None) value
                text = "This figure could not be analyzed and was skipped."

            doc.add_heading(f"Figure {idx}", level=2)
            doc.add_paragraph(text)

    # --- Footer note ---
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        "Generated by the Research Reconstructor pipeline "
        "(local Ollama graph extraction + Groq vision & synthesis)."
    )
    footer.runs[0].italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------------
def main():
    st.set_page_config(page_title="Research Reconstructor", page_icon="🧠", layout="centered")

    st.title("🧠 Research Reconstructor")
    st.caption(
        "Upload a research paper. We'll extract its structure and charts, "
        "build a knowledge graph, and synthesize a polished executive summary."
    )

    uploaded_file = st.file_uploader("Upload a research paper (PDF)", type=["pdf"])

    if uploaded_file is None:
        st.info("Upload a PDF to get started.")
        return

    if st.button("Run Reconstruction Pipeline", type="primary"):
        pdf_bytes = uploaded_file.read()

        with st.status("Running pipeline...", expanded=True) as status:
            # --- 1. Text extraction & chunking ---
            st.write("📄 Extracting text from PDF...")
            full_text = extract_text_from_pdf(pdf_bytes)
            chunks = chunk_text(full_text)
            st.write(f"Split into {len(chunks)} text chunk(s).")

            # --- 2. Local graph extraction (Ollama) + push to Neo4j Aura ---
            st.write("🕸️ Extracting knowledge graph via local Ollama (llama3.1:8b)...")
            try:
                store = GraphStore()
                graph_documents = store.run_pipeline(chunks)
                graph_summary = summarize_graph_documents(graph_documents)
                store.close()
            except Exception as exc:
                logger.error("Graph extraction/storage failed: %s", exc)
                graph_summary = f"Graph extraction step failed: {exc}"
                st.warning("Graph extraction failed — continuing without graph insights.")

            # --- 3. Chart extraction & Groq vision analysis ---
            st.write("📊 Extracting and analyzing charts with Groq vision...")
            # Write PDF to a temp path for PyMuPDF-based chart extraction.
            # Uses tempfile.gettempdir() (NOT a hardcoded "/tmp") so this
            # works cross-platform, including on Windows where "/tmp"
            # doesn't exist.
            temp_pdf_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
            with open(temp_pdf_path, "wb") as f:
                f.write(pdf_bytes)

            chart_images = extract_images_from_pdf(temp_pdf_path)
            chart_analyses = []
            for i, image in enumerate(chart_images):
                st.write(f"  Analyzing figure {i + 1}/{len(chart_images)}...")
                result = analyze_chart(image, context_hint=full_text[:500])
                chart_analyses.append(result)

            # --- 4. Report synthesis (Groq) ---
            st.write("✍️ Synthesizing executive summary with Groq...")
            client = get_groq_client()
            try:
                summary_text = synthesize_report(
                    client=client,
                    paper_text_excerpt=full_text,
                    chart_analyses=chart_analyses,
                    graph_summary=graph_summary,
                )
            except RuntimeError as exc:
                status.update(label="Pipeline failed", state="error")
                st.error(f"❌ {exc}")
                st.stop()

            # --- 5. Compile .docx ---
            st.write("📝 Compiling Word document...")
            report_title = os.path.splitext(uploaded_file.name)[0] + " — Executive Summary"
            docx_buffer = compile_docx(
                title=report_title,
                summary_text=summary_text,
                graph_summary=graph_summary,
                chart_analyses=chart_analyses,
            )

            status.update(label="Pipeline complete!", state="complete")

        st.success("Your executive summary is ready.")
        st.download_button(
            label="⬇️ Download Executive Summary (.docx)",
            data=docx_buffer,
            file_name=f"{os.path.splitext(uploaded_file.name)[0]}_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        with st.expander("Preview: Executive Summary Text"):
            st.write(summary_text)

        with st.expander("Preview: Knowledge Graph Summary"):
            st.write(graph_summary)


if __name__ == "__main__":
    main()